import docx
import datetime
from sys import argv
from random import randint
from months import Months
from num2words import num2words
import os

doc = docx.Document("123.docx")
mobyli_bill = str(argv[1])
Enet_bill = str(argv[2])
now = datetime.datetime.now()
init_str = doc.tables[1].rows[0].cells[0].text
doc.tables[1].rows[0].cells[0].text = str(init_str[:7] + str(randint(0, 1000)) + " " +
                                          init_str[11:14] + " " + format(now.day, '02d') + "\t" +
                                          str(str(Months(now.month).name)) + "\t" + str(now.year) + "г.")
init_str = doc.tables[2].rows[2].cells[1].text
doc.tables[2].rows[2].cells[1].text = str(init_str[:1] + " " + str(randint(0, 100)) + " " + init_str[3: 6] +
                                          format(now.day, '02d') + "." + format(now.month, '02d') + "." + str(now.year))

doc.tables[3].rows[1].cells[4].text = mobyli_bill
doc.tables[3].rows[1].cells[5].text = mobyli_bill
doc.tables[3].rows[2].cells[4].text = Enet_bill
doc.tables[3].rows[2].cells[5].text = Enet_bill
doc.tables[4].rows[0].cells[1].text = str("{:.2f}".format(float(mobyli_bill) + float(Enet_bill)))
doc.tables[4].rows[1].cells[1].text = str("{:.2f}".format(float(0.2 * (int(mobyli_bill) + int(Enet_bill)))))
doc.tables[4].rows[2].cells[1].text = str("{:.2f}".format(float(mobyli_bill) + float(Enet_bill)))

init_str = doc.tables[4].rows[3].cells[0].text
doc.tables[4].rows[3].cells[0].text = str(init_str[:20] + "2," + init_str[22:32] +
                                          str("{:.2f}".format(float(mobyli_bill) + float(Enet_bill))) + " руб.")

choose = " рублей "
init_str = doc.tables[4].rows[4].cells[1].text
doc.tables[4].rows[4].cells[1].text = str(num2words(int(float(mobyli_bill) + float(Enet_bill)), lang="ru").capitalize()
                                           + choose + format(int(100*(float(mobyli_bill) + float(Enet_bill)))%100, '02d')
                                          + " копеек" + init_str[14:])

doc.save("321.docx")
os.system("abiword --to=pdf 321.docx")
